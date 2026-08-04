# Copyright (c) 2026, Tridz Technologies Pvt Ltd and contributors
# For license information, please see license.txt

"""Convert the HTML produced by ``render_document_html`` into a .docx file.

This is the second stage of the markdown -> {PDF, DOCX} pipeline. The HTML
contract is defined and enforced by ``huf.ai.artifacts.render.html`` — this
module does not accept arbitrary HTML, only the specific tag/class
vocabulary that ``render_document_html`` produces:

- Block tags: p, h1-h6, ul, ol, li, table, thead, tbody, tr, th, td,
  blockquote, hr, pre
- Inline tags: strong, em, a, img, br, code, span, div
- Alignment via class="text-left|text-center|text-right" on a block element
- Indent via class="indent-1|indent-2|indent-3" on a block element
- Document components (doc-header, callout, metric-grid, split, data-table,
  etc.) via a class from ``huf.ai.artifacts.render.components.COMPONENTS`` —
  see the "Component dispatch" section below.

The HTML is walked with the stdlib ``html.parser.HTMLParser`` (not
BeautifulSoup) to keep this module dependency-free beyond python-docx.

Component dispatch
-------------------
``COMPONENTS`` (huf/ai/artifacts/render/components.py) is the single source
of truth for what a component class means in DOCX terms — every entry's
"docx" key is a data recipe, not code, so this module is an INTERPRETER for
those recipes rather than a second definition of what "doc-header" or
"callout" looks like. Rendering functions below accept a ``container``
(either the top-level ``Document`` or a table ``_Cell``) rather than always
a ``Document``, since python-docx's ``Document`` and ``_Cell`` both expose
``add_paragraph``/``add_table`` — this lets component recipes nest (e.g. a
callout's children, or a split cell's content) using the exact same
rendering functions as the top level.
"""

import base64
import io
import re
from html.parser import HTMLParser

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from huf.ai.artifacts.render.components import COMPONENTS, COMPONENT_CLASSES, THEME, resolve_theme_token


#: Maps a block element's alignment class to a python-docx alignment constant.
_ALIGNMENTS = {
	"text-left": WD_ALIGN_PARAGRAPH.LEFT,
	"text-center": WD_ALIGN_PARAGRAPH.CENTER,
	"text-right": WD_ALIGN_PARAGRAPH.RIGHT,
}

#: Maps a block element's indent class to an indent level (in half-inches).
_INDENTS = {
	"indent-1": 1,
	"indent-2": 2,
	"indent-3": 3,
}

#: Block-level tags that terminate the "current" node when closed. Everything
#: else (inline tags, and stray tags outside the sanitized vocabulary) is
#: treated as text-bearing but does not open a new node in the tree, UNLESS
#: it carries a registry component class — see ``_opens_node``.
_BLOCK_TAGS = {
	"p", "h1", "h2", "h3", "h4", "h5", "h6",
	"ul", "ol", "li",
	"table", "thead", "tbody", "tr", "th", "td",
	"blockquote", "hr", "pre",
	"div",
}

#: Container tags whose children must each render as their own separate docx
#: block, never flattened into a single joined string via full_text. div is
#: only ever used by render_document_html for a columns-N section wrapper (or
#: a document component) - a structural container, not a text-bearing leaf
#: like p/blockquote.
_CONTAINER_TAGS = ("table", "thead", "tbody", "tr", "th", "td", "div")

#: Tags whose text (and descendant text) should be dropped from the parent's
#: rendered text rather than concatenated inline. None of the tags in the
#: sanitized vocabulary need this, but "script"/"style" are excluded
#: defensively in case they ever slip through. "style" text is not simply
#: discarded, though - see _BodyTreeBuilder.style_text_parts - an author's
#: <style> block is the only place a re-themed document's colours live, and a
#: document that redefined --accent etc. in a <style> block used to export
#: with the registry's hardcoded default palette because this text never
#: reached anything past the parser. Keeping it OUT of the body (so it never
#: leaks as literal paragraph text - that part already worked) while ALSO
#: capturing it separately is what lets html_to_docx() parse the author's
#: :root overrides and recolour every recipe to match.
_IGNORED_TEXT_TAGS = {"script", "style"}


class _Node:
	"""A single element in the simplified DOM tree built from the HTML body."""

	def __init__(self, tag: str, attrs: dict[str, str], parent: "_Node | None" = None):
		self.tag = tag
		self.attrs = attrs
		self.parent = parent
		self.children: list["_Node"] = []
		#: Accumulated text for this node, used for leaf-like elements (p, h1-h6,
		#: li, th, td, blockquote). Container elements (table, tr, ul, ol) only
		#: use ``children``.
		self.text_parts: list[str] = []
		#: Text and child nodes in the exact order the parser encountered
		#: them - ("text", str) or ("node", _Node). text_parts/children above
		#: are each flattened views of this (kept for the existing code that
		#: only needs one or the other); this is what _all_descendant_text
		#: needs to reconstruct real document order across a mix of bare text
		#: and nested elements, e.g. "<strong>Label:</strong> tail text".
		self._ordered_parts: list[tuple[str, "str | _Node"]] = []

	def add_text(self, data: str) -> None:
		self.text_parts.append(data)
		self._ordered_parts.append(("text", data))

	def add_child(self, child: "_Node") -> None:
		self.children.append(child)
		self._ordered_parts.append(("node", child))

	@property
	def text(self) -> str:
		return "".join(self.text_parts).strip()

	def classes(self) -> set[str]:
		class_attr = self.attrs.get("class") or ""
		return set(class_attr.split())

	@property
	def full_text(self) -> str:
		"""Own text plus the text of any nested block children, joined by
		newlines. Markdown wraps blockquote (and sometimes list item) content
		in a nested ``<p>``, so a node's directly-owned text alone can miss
		that content — this walks block descendants to recover it.
		"""
		parts = [self.text] if self.text else []
		for child in self.children:
			if child.tag in _BLOCK_TAGS and child.tag not in _CONTAINER_TAGS:
				child_text = child.full_text
				if child_text:
					parts.append(child_text)
		return "\n".join(parts)


def _all_descendant_text(node: "_Node") -> str:
	"""Every bit of text under ``node``, in true document order, regardless
	of tag - unlike ``full_text`` (which only walks ``_BLOCK_TAGS`` minus
	``_CONTAINER_TAGS``, deliberately excluding div/table/td/etc.), this
	walks EVERY child node.

	Needed because ``div`` is both a ``_BLOCK_TAG`` and a ``_CONTAINER_TAG``:
	a value wrapped in a child div - e.g.
	``<div class="metric" data-label="X"><div class="metric-value">$12.5M</div></div>``
	- puts "$12.5M" on the *child* node, and ``full_text`` skips container
	children on purpose, so the value silently disappeared even though
	``full_text`` looked like the obvious thing to reach for. Observed in
	production: a metric card exported with its label but no value.

	Walks ``_ordered_parts`` rather than ``text_parts``/``children``
	separately so text that comes before/after a nested element - e.g.
	``<strong>Label:</strong> tail text`` - is reassembled in the order it
	was written rather than all-text-then-all-children.
	"""
	segments = []
	for kind, value in node._ordered_parts:
		if kind == "text":
			segments.append(value)
		elif value.tag != "img":
			segments.append(_all_descendant_text(value))
	return "".join(segments).strip()


def _opens_node(tag: str, attrs: dict[str, str]) -> bool:
	"""Whether a start tag should open a new ``_Node`` and become the
	parser's "current" element.

	True for the fixed block-tag vocabulary (unchanged from before), and
	ALSO true for any tag carrying a registry component class (e.g. a
	``<span class="brand">``) — components can be authored on inline tags
	that would otherwise just fold their text into the enclosing block, and
	the DOCX walker needs a real node to dispatch the recipe against. This
	reads from ``COMPONENT_CLASSES`` rather than hardcoding tag names, so a
	future component class works regardless of which tag it's applied to.
	"""
	if tag in _BLOCK_TAGS:
		return True
	classes = set((attrs.get("class") or "").split())
	return bool(classes & COMPONENT_CLASSES)


class _BodyTreeBuilder(HTMLParser):
	"""Builds a simplified tree of ``_Node`` objects from the sanitized HTML body.

	Only tags reach this parser that were already allowed through
	``render_document_html``'s bleach sanitization, so no further validation
	of the tag/attribute vocabulary is performed here.
	"""

	def __init__(self):
		super().__init__(convert_charrefs=True)
		self.root = _Node("root", {})
		self._current = self.root
		self._ignored_depth = 0
		#: Which _IGNORED_TEXT_TAGS tag is currently suppressing body text, as
		#: a stack rather than a single name - handle_endtag needs to know
		#: which tag is closing even though script/style never nest in
		#: practice, so a stack costs nothing and avoids a latent bug if they
		#: ever did.
		self._ignored_tag_stack: list[str] = []
		#: Raw text found inside <style> tags, kept separate from the node
		#: tree entirely (see _IGNORED_TEXT_TAGS) so html_to_docx() can parse
		#: an author's :root overrides out of it after the walk finishes.
		self.style_text_parts: list[str] = []

	def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
		if tag in _IGNORED_TEXT_TAGS:
			self._ignored_depth += 1
			self._ignored_tag_stack.append(tag)
			return
		if self._ignored_depth:
			return

		attr_dict = {key: (value or "") for key, value in attrs}

		if _opens_node(tag, attr_dict):
			node = _Node(tag, attr_dict, parent=self._current)
			self._current.add_child(node)
			self._current = node
		elif tag == "img":
			# img is a void element (no matching handle_endtag) — record it as
			# a leaf child of the current node so text order is preserved.
			node = _Node(tag, attr_dict, parent=self._current)
			self._current.add_child(node)
		elif tag == "br":
			self._current.add_text("\n")
		# Other inline tags (strong, em, a, span, code without a component
		# class) don't need their own node — their text is folded into the
		# enclosing block's text.

	def handle_startendtag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
		if tag == "img":
			self.handle_starttag(tag, attrs)
		elif tag == "br":
			self.handle_starttag(tag, attrs)
		elif tag == "hr":
			self.handle_starttag(tag, attrs)
			self.handle_endtag(tag)

	def handle_endtag(self, tag: str) -> None:
		if tag in _IGNORED_TEXT_TAGS:
			self._ignored_depth = max(0, self._ignored_depth - 1)
			if self._ignored_tag_stack:
				self._ignored_tag_stack.pop()
			return
		if self._ignored_depth:
			return

		# self._current.tag can only equal `tag` here if a node was actually
		# opened for it (see _opens_node) — a plain <span> or <strong> never
		# becomes `self._current`, so its close tag is a no-op, same as before.
		if self._current is not self.root and self._current.tag == tag:
			self._current = self._current.parent or self.root

	def handle_data(self, data: str) -> None:
		if self._ignored_depth:
			# Still inside script/style - the text must not join the node
			# tree (that's what keeps CSS out of the body), but style text
			# specifically is captured for theme parsing. script text is
			# genuinely dropped; nothing downstream needs it.
			if self._ignored_tag_stack and self._ignored_tag_stack[-1] == "style":
				self.style_text_parts.append(data)
			return
		self._current.add_text(data)


#: Matches one ``:root { ... }`` block's body. Author CSS may contain more
#: than one :root rule (e.g. a base one plus a media-query-guarded override
#: that still landed unconditionally after sanitization) - finditer() below
#: walks them in source order so a later block's declarations win, matching
#: normal CSS cascade behaviour for repeated declarations of the same
#: property.
_ROOT_BLOCK_RE = re.compile(r":root\s*\{([^}]*)\}", re.DOTALL)

#: Matches one ``--token: value;`` custom-property declaration inside a
#: :root block. Deliberately permissive about the token charset and
#: whitespace/newlines - author CSS is free-form text we don't control.
_CSS_CUSTOM_PROP_RE = re.compile(r"--([A-Za-z0-9_-]+)\s*:\s*([^;}]+);?")

#: Matches a bare #RGB or #RRGGBB colour literal, nothing else. Used to
#: reject non-colour custom properties (e.g. ``--gap: 8pt``) before they can
#: reach OOXML, which has no notion of any other CSS value shape.
_HEX_COLOR_RE = re.compile(r"^#([0-9A-Fa-f]{3}|[0-9A-Fa-f]{6})$")


def _normalize_hex_color(value: str) -> str | None:
	"""Return a ``#RRGGBB`` literal for a ``#RGB``/``#RRGGBB`` value, or
	``None`` if ``value`` isn't a hex colour at all (e.g. ``8pt``, ``red``,
	``rgb(1,2,3)``). OOXML colour attributes require exactly 6 hex digits -
	python-docx's ``RGBColor.from_string`` does not expand the 3-digit CSS
	shorthand, so ``#D33`` reaching python-docx unexpanded raises instead of
	degrading gracefully.
	"""
	match = _HEX_COLOR_RE.match(value.strip())
	if not match:
		return None
	digits = match.group(1)
	if len(digits) == 3:
		digits = "".join(ch * 2 for ch in digits)
	return f"#{digits.upper()}"


def _extract_theme(style_text: str) -> dict:
	"""Parse an author's ``:root { --token: value; }`` overrides out of the
	raw text captured from every <style> tag, and return the effective theme:
	the registry defaults (``components.THEME``) updated with whatever the
	author validly overrode.

	This must never raise - a malformed or hostile <style> block should
	degrade to the default palette, not break the export. Every failure mode
	(no :root block, an unparsable declaration, a non-colour value, an
	unknown token) is handled by simply not updating that entry, rather than
	by catching exceptions from deep inside a parse - the regexes above are
	permissive enough that a try/except around the whole function is only a
	last-resort safety net, not the primary defence.
	"""
	theme = dict(THEME)
	if not style_text:
		return theme

	try:
		for root_match in _ROOT_BLOCK_RE.finditer(style_text):
			block_body = root_match.group(1)
			for decl_match in _CSS_CUSTOM_PROP_RE.finditer(block_body):
				token = decl_match.group(1).strip()
				raw_value = decl_match.group(2).strip()
				normalized = _normalize_hex_color(raw_value)
				if normalized is None:
					# Non-colour custom property (e.g. --gap: 8pt) or a CSS
					# colour keyword/function we don't attempt to resolve -
					# leave whatever this token was already set to (default,
					# or an earlier :root block's value) rather than write a
					# value OOXML can't use.
					continue
				theme[token] = normalized
	except Exception:
		# Never let a theme-parsing surprise take down the whole export -
		# fall back to the registry defaults, exactly as if the document
		# carried no <style> block at all.
		return dict(THEME)

	return theme


def _apply_alignment_and_indent(paragraph, node: _Node) -> None:
	classes = node.classes()

	for class_name, alignment in _ALIGNMENTS.items():
		if class_name in classes:
			paragraph.alignment = alignment
			break

	for class_name, level in _INDENTS.items():
		if class_name in classes:
			paragraph.paragraph_format.left_indent = Inches(0.5 * level)
			break


def _add_heading(container, text: str, level: int):
	"""Add a heading paragraph to ``container``. Used both for plain h1-h6
	tags and the "doc-title" component recipe. ``Document.add_heading``
	isn't available on a table ``_Cell``, so this applies the equivalent
	built-in heading/title style directly via ``add_paragraph`` instead,
	which both ``Document`` and ``_Cell`` support.
	"""
	style = f"Heading {level}" if level else "Title"
	return container.add_paragraph(text, style=style)


def _add_image(container, node: _Node) -> None:
	src = node.attrs.get("src") or ""
	if not src.startswith("data:"):
		# http(s) URLs (or anything else) are skipped — no network calls.
		return

	try:
		header, _, encoded = src.partition(",")
		if ";base64" not in header:
			return
		image_bytes = base64.b64decode(encoded)
		# Document.add_picture() is a convenience method not available on a
		# table _Cell — adding the run explicitly works in either container.
		paragraph = container.add_paragraph()
		paragraph.add_run().add_picture(io.BytesIO(image_bytes))
	except Exception:
		# Best-effort: a malformed data URI should not crash the export.
		pass


def _add_table(container, table_node: _Node, theme: dict):
	"""Build a plain docx table from a sanitized <table> node's rows/cells.

	Returns the created ``Table`` (or ``None`` if the source had no rows) so
	callers - notably the "data-table" component recipe - can post-process
	it (e.g. shade the header row) without re-walking the HTML.
	"""
	rows: list[list[_Node]] = []
	for section in table_node.children:
		if section.tag in ("thead", "tbody"):
			for tr in section.children:
				if tr.tag == "tr":
					rows.append([cell for cell in tr.children if cell.tag in ("th", "td")])
		elif section.tag == "tr":
			rows.append([cell for cell in section.children if cell.tag in ("th", "td")])

	if not rows:
		return None

	num_cols = max(len(row) for row in rows)
	table = container.add_table(rows=len(rows), cols=num_cols)
	table.style = "Table Grid"

	for row_index, row in enumerate(rows):
		for col_index, cell_node in enumerate(row):
			_fill_plain_table_cell(table.cell(row_index, col_index), cell_node, theme)

	return table


def _fill_plain_table_cell(cell, node: _Node, theme: dict) -> None:
	"""Fill one <td>/<th> cell of a plain sanitized <table> (data-table).

	A cell that's just text (``<td><strong>Digital Omnichannel</strong></td>``)
	takes the fast ``cell.text =`` path unchanged from before. A cell that
	contains a component element - the only real case is
	``<td><span class="status-badge ...">In Progress</span></td>`` - used to
	go through the same ``full_text`` path, which drops it: ``span`` is not
	in ``_BLOCK_TAGS``, so a status-badge span's text was never included in
	the td's ``full_text`` even though the span DOES open its own ``_Node``
	(component classes force that - see ``_opens_node``). The badge's text
	vanished from the export entirely, and even where a future component
	happened to still have SOME text reach the cell, it would arrive as bare
	text with the recipe's styling (shading/bold/etc.) never applied, since
	nothing here ever called `_render_node` for cell content. Routing
	children through `_render_node` when there are any is what lets a
	table-cell badge actually dispatch to its "run" recipe.
	"""
	if node.children:
		_clear_cell(cell)
		for child in node.children:
			_render_node(cell, child, theme)
	else:
		cell.text = node.full_text


def _add_list(container, list_node: _Node, style: str) -> None:
	for item in list_node.children:
		if item.tag == "li":
			container.add_paragraph(item.full_text, style=style)


def _render_node(container, node: _Node, theme: dict) -> None:
	class_name, recipe = _component_recipe(node)
	if recipe is not None:
		_render_component(container, node, class_name, recipe, theme)
		return

	tag = node.tag

	if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
		level = int(tag[1])
		_add_heading(container, node.full_text, level)
	elif tag == "p":
		if node.children and all(child.tag == "img" for child in node.children) and not node.full_text:
			for child in node.children:
				_add_image(container, child)
			return
		paragraph = container.add_paragraph(node.full_text)
		_apply_alignment_and_indent(paragraph, node)
	elif tag == "blockquote":
		paragraph = container.add_paragraph(node.full_text, style="Intense Quote")
		_apply_alignment_and_indent(paragraph, node)
	elif tag == "ul":
		_add_list(container, node, "List Bullet")
	elif tag == "ol":
		_add_list(container, node, "List Number")
	elif tag == "table":
		_add_table(container, node, theme)
	elif tag == "pre":
		container.add_paragraph(node.full_text)
	elif tag == "hr":
		container.add_paragraph("_" * 40)
	elif tag == "img":
		_add_image(container, node)
	elif tag == "div":
		_render_div(container, node, theme)
	elif node.children:
		# Unclassed/unknown container tag - degrade gracefully by rendering
		# children in document order rather than dropping them.
		for child in node.children:
			_render_node(container, child, theme)
	# thead/tbody/tr/th/td/li are handled by their container (table/ul/ol)
	# and should never appear as direct children of the root.


_COLUMNS_CLASS_RE = re.compile(r"^columns-([23])$")


def _render_div(container, node: _Node, theme: dict) -> None:
	"""Render a <div> that is not carrying a registry component class —
	currently either a ``columns-N`` section wrapper, or an unclassed
	passthrough container. A div without a recognised columns class renders
	as a plain passthrough (its children rendered inline, no section
	change), so future non-columns div usage degrades gracefully rather than
	silently dropping content.
	"""
	column_count = None
	for class_name in node.classes():
		match = _COLUMNS_CLASS_RE.match(class_name)
		if match:
			column_count = int(match.group(1))
			break

	# Columns require a real docx section, which only exists at the
	# Document level - a div nested inside a table cell (e.g. inside a
	# callout) can't open one, so it degrades to a plain passthrough too.
	if column_count is None or not hasattr(container, "add_section"):
		for child in node.children:
			_render_node(container, child, theme)
		return

	_set_section_columns(container.add_section(WD_SECTION.CONTINUOUS), column_count)
	for child in node.children:
		_render_node(container, child, theme)
	# Close the columns region: a fresh CONTINUOUS section reset to a single
	# column, so content after the div returns to normal single-column flow
	# rather than inheriting the multi-column layout indefinitely.
	_set_section_columns(container.add_section(WD_SECTION.CONTINUOUS), 1)


def _set_section_columns(section, column_count: int) -> None:
	"""Set the number of layout columns on a docx section via its raw
	sectPr XML — python-docx has no public column-count API. A freshly
	created section's sectPr has no <w:cols> element yet, so one is created
	if absent rather than assumed to already exist.
	"""
	sect_pr = section._sectPr
	cols = sect_pr.find(qn("w:cols"))
	if cols is None:
		cols = OxmlElement("w:cols")
		sect_pr.append(cols)
	cols.set(qn("w:num"), str(column_count))


# ---------------------------------------------------------------------------
# Component dispatch
#
# The functions below interpret a component's "docx" recipe (read from
# COMPONENTS at call time - never copied/hardcoded here) against the node it
# was found on. Each one accepts the same `container` (Document or _Cell)
# that ordinary node rendering uses, so nested content - a callout's
# children, a split cell's paragraphs, a metric's value/label - keeps
# flowing through the same _render_node() machinery as everything else.
# ---------------------------------------------------------------------------


def _component_recipe(node: _Node):
	"""Return (class_name, docx recipe dict) for the first registry class
	found on `node`, or (None, None) if it carries none. Iterates
	COMPONENTS (dict order) rather than the node's own (unordered) class set
	so dispatch is deterministic if an element ever carried more than one
	registry class.
	"""
	classes = node.classes()
	for class_name in COMPONENTS:
		if class_name in classes:
			return class_name, COMPONENTS[class_name]["docx"]
	return None, None


def _render_component(container, node: _Node, class_name: str, recipe: dict, theme: dict) -> None:
	kind = recipe.get("type")
	if kind == "table":
		_render_table_recipe(container, node, recipe, theme)
	elif kind == "table_row_of_cells":
		_render_metric_grid(container, node, recipe, theme)
	elif kind == "cell":
		_render_cell_recipe(container, node, recipe, theme)
	elif kind == "run":
		_render_run_recipe(container, node, recipe, theme)
	elif kind == "paragraph":
		_render_paragraph_recipe(container, node, recipe, theme)
	elif kind == "heading":
		_render_heading_recipe(container, node, recipe, theme)
	elif kind == "page_break":
		_render_page_break_recipe(container, node, recipe)
	elif kind == "footer":
		_render_footer_recipe(container, node, recipe, theme)
	else:
		# Unknown recipe type: degrade gracefully rather than drop content.
		for child in node.children:
			_render_node(container, child, theme)


def _render_table_recipe(container, node: _Node, recipe: dict, theme: dict) -> None:
	"""type: "table". Two shapes, distinguished by the recipe itself:

	- Has "cols": the class is on an arbitrary container (div) whose direct
	  children become the table's cells (doc-header, callout, split).
	- No "cols": the class is on a real sanitized <table> (data-table) - the
	  existing row/cell walker builds the table, this only adds header
	  shading on top.
	"""
	if "cols" in recipe:
		_render_synthetic_table(container, node, recipe, theme)
	else:
		_render_data_table(container, node, recipe, theme)


def _render_data_table(container, node: _Node, recipe: dict, theme: dict) -> None:
	table = _add_table(container, node, theme)
	if table is None or not table.rows:
		return

	header_shading = resolve_theme_token(recipe.get("header_shading"), theme) if recipe.get("header_shading") else None
	header_color = resolve_theme_token(recipe.get("header_color"), theme) if recipe.get("header_color") else None
	for cell in table.rows[0].cells:
		if header_shading:
			_shade_cell(cell, header_shading)
		if header_color:
			for paragraph in cell.paragraphs:
				for run in paragraph.runs:
					run.font.color.rgb = _rgb_color(header_color)
					run.font.bold = True


def _render_synthetic_table(container, node: _Node, recipe: dict, theme: dict) -> None:
	cols = recipe["cols"]

	# Assign each direct child its own cell in document order; if there are
	# more children than columns (or exactly one column, e.g. callout), the
	# overflow lands in the last cell together rather than being dropped.
	cell_children: list[list[_Node]] = [[] for _ in range(cols)]
	for index, child in enumerate(node.children):
		cell_children[min(index, cols - 1)].append(child)

	table = container.add_table(rows=1, cols=cols)
	_set_table_borders(table, recipe.get("borders", True))

	widths = recipe.get("widths")
	if widths:
		_set_column_widths(table, widths)

	shading = recipe.get("shading")
	col_align = recipe.get("col_align")

	for col_index in range(cols):
		cell = table.cell(0, col_index)
		if shading and cols == 1:
			_shade_cell(cell, resolve_theme_token(shading, theme))

		contents = cell_children[col_index]
		if contents:
			_clear_cell(cell)
			for child in contents:
				_render_node(cell, child, theme)
		elif cols == 1 and node.text:
			# A node like <div class="callout"><strong>X:</strong> tail
			# text</div> has no child _Node at all (see _opens_node - a
			# bare <strong> without a component class never opens one), so
			# ALL of its text - "X: tail text" - sits on the *node's own*
			# text_parts, and `contents` above is empty. Observed in
			# production: an executive-summary callout authored exactly
			# this way exported with a visually correct box and NO text in
			# it. Gated to cols == 1 (callout is the only single-column
			# user of this recipe shape): for a multi-column recipe
			# (doc-header, split) the parent node's text can't be
			# attributed to one particular empty column, so this fallback
			# would duplicate it into every empty column instead of fixing
			# anything. _render_cell_recipe has the equivalent fallback for
			# the split/metric "cell" recipe shape; this mirrors it for the
			# "table" shape.
			_clear_cell(cell)
			cell.add_paragraph(_all_descendant_text(node))

		if col_align and col_index < len(col_align):
			alignment = _ALIGNMENTS.get(f"text-{col_align[col_index]}")
			if alignment is not None:
				for paragraph in cell.paragraphs:
					paragraph.alignment = alignment


def _render_metric_grid(container, node: _Node, recipe: dict, theme: dict) -> None:
	"""type: "table_row_of_cells" (metric-grid). One table, 2 columns per
	row, so 4 .metric children become a 2x2 table - matching how the PDF
	lays them out via CSS grid. Each metric child is rendered through the
	normal component dispatch (its own "metric" recipe fills the cell we
	hand it), so this function only owns the row/column bookkeeping.
	"""
	metrics = node.children
	if not metrics:
		return

	cols = 2
	rows = (len(metrics) + cols - 1) // cols
	table = container.add_table(rows=rows, cols=cols)

	for index, metric_node in enumerate(metrics):
		row_index, col_index = divmod(index, cols)
		cell = table.cell(row_index, col_index)
		# Cell clearing is owned by whichever component recipe fills the
		# cell (see _render_cell_recipe/_render_metric_value_cell) - not
		# done here too, since _clear_cell is only safe to call once on a
		# cell that still has its default empty paragraph.
		_render_node(cell, metric_node, theme)


def _render_cell_recipe(container, node: _Node, recipe: dict, theme: dict) -> None:
	"""type: "cell". Fills `container` (an already-created table cell) with
	this node's content - it does NOT create a new cell itself, since the
	caller (a "table"/"table_row_of_cells" recipe) already built the
	physical cell and handed it in as `container`.

	Two shapes, distinguished by the recipe's own keys: a plain cell
	(split-main/split-side - optional shading, children render normally) or
	a metric cell (has "value_size_pt" - node.text becomes a large bold
	value run, node.attrs[label_from] becomes a small label paragraph).
	"""
	if "value_size_pt" in recipe:
		_render_metric_value_cell(container, node, recipe, theme)
		return

	shading = recipe.get("shading")
	if shading:
		_shade_cell(container, resolve_theme_token(shading, theme))

	if node.children:
		_clear_cell(container)
		for child in node.children:
			_render_node(container, child, theme)
	elif node.text:
		_clear_cell(container)
		container.add_paragraph(node.text)


def _render_metric_value_cell(container, node: _Node, recipe: dict, theme: dict) -> None:
	shading = recipe.get("shading")
	if shading:
		_shade_cell(container, resolve_theme_token(shading, theme))

	_clear_cell(container)

	# node.text (not _all_descendant_text) would only see text sitting
	# DIRECTLY on the .metric node - an author who wraps the value in its
	# own <div class="metric-value"> (matching the CSS, which styles
	# .metric-value separately from .metric) puts that text on a CHILD
	# node instead. node.full_text doesn't help either: div is a
	# _CONTAINER_TAG, so full_text's block-descendant walk skips it on
	# purpose. Observed in production: the metric card rendered with its
	# label but a blank value. _all_descendant_text has no such exclusion.
	value_text = _all_descendant_text(node)
	value_paragraph = container.add_paragraph()
	value_run = value_paragraph.add_run(value_text)
	value_run.font.bold = True
	value_run.font.size = Pt(recipe["value_size_pt"])
	value_color = recipe.get("value_color")
	if value_color:
		value_run.font.color.rgb = _rgb_color(resolve_theme_token(value_color, theme))

	label_attr = recipe.get("label_from")
	label_text = node.attrs.get(label_attr, "") if label_attr else ""
	if label_text:
		label_paragraph = container.add_paragraph()
		label_run = label_paragraph.add_run(label_text)
		label_run.font.size = Pt(recipe["label_size_pt"])
		label_color = recipe.get("label_color")
		if label_color:
			label_run.font.color.rgb = _rgb_color(resolve_theme_token(label_color, theme))


def _render_run_recipe(container, node: _Node, recipe: dict, theme: dict) -> None:
	"""type: "run" (brand, status-badge). A single run in its own paragraph,
	optionally bold/coloured/shaded/sized.
	"""
	paragraph = container.add_paragraph()
	run = paragraph.add_run(node.text)
	if recipe.get("bold"):
		run.font.bold = True
	color = recipe.get("color")
	if color:
		run.font.color.rgb = _rgb_color(resolve_theme_token(color, theme))
	size_pt = recipe.get("size_pt")
	if size_pt:
		run.font.size = Pt(size_pt)
	shading = recipe.get("shading")
	if shading:
		_shade_run(run, resolve_theme_token(shading, theme))


def _render_paragraph_recipe(container, node: _Node, recipe: dict, theme: dict) -> None:
	"""type: "paragraph" (doc-meta, doc-subtitle)."""
	paragraph = container.add_paragraph(node.full_text)

	align = recipe.get("align")
	alignment = _ALIGNMENTS.get(f"text-{align}") if align else None
	if alignment is not None:
		paragraph.alignment = alignment

	size_pt = recipe.get("size_pt")
	color = recipe.get("color")
	if paragraph.runs and (size_pt or color):
		run = paragraph.runs[0]
		if size_pt:
			run.font.size = Pt(size_pt)
		if color:
			run.font.color.rgb = _rgb_color(resolve_theme_token(color, theme))


def _render_heading_recipe(container, node: _Node, recipe: dict, theme: dict) -> None:
	"""type: "heading" (doc-title)."""
	_add_heading(container, node.full_text, recipe.get("level", 1))


def _render_page_break_recipe(container, node: _Node, recipe: dict) -> None:
	"""type: "page_break" (page-break). A real Word page break, not a
	visible rule - see components.py's "page-break" comment: an
	agent-authored dashed-rule version of this got WeasyPrint to paint a
	stray line above every forced break. This only forces the break.

	``add_break(WD_BREAK.PAGE)`` requires an existing run to carry the
	break character, and is not available on a bare paragraph - a fresh run
	on a fresh paragraph is the standard python-docx idiom for an
	unaccompanied page break.
	"""
	paragraph = container.add_paragraph()
	paragraph.add_run().add_break(WD_BREAK.PAGE)


def _render_footer_recipe(container, node: _Node, recipe: dict, theme: dict) -> None:
	"""type: "footer" (doc-footer). Writes into the document's real Word
	section footer, so the text repeats on every page - see components.py's
	"doc-footer" comment: a plain body paragraph only ever appeared once,
	mid-flow, defeating the point of a running footer.

	Only meaningful at the Document level: `container` here is always
	`doc` in practice (doc-footer is never nested inside another
	component's cell in the sanitized vocabulary), but this checks for
	`sections` defensively rather than assuming it, since a table _Cell has
	no notion of a footer to write into.
	"""
	if not hasattr(container, "sections"):
		# No Document reachable from here (e.g. doc-footer nested inside a
		# table cell) - degrade to an ordinary paragraph in place rather
		# than silently dropping the footer text.
		container.add_paragraph(node.full_text)
		return

	footer = container.sections[0].footer
	footer.is_linked_to_previous = False
	paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
	paragraph.text = ""
	run = paragraph.add_run(node.full_text)

	size_pt = recipe.get("size_pt")
	if size_pt:
		run.font.size = Pt(size_pt)
	color = recipe.get("color")
	if color:
		run.font.color.rgb = _rgb_color(resolve_theme_token(color, theme))

	_add_page_number_field(paragraph)


def _add_page_number_field(paragraph) -> None:
	"""Append " | Page {PAGE} of {NUMPAGES}" to a footer paragraph as real
	Word fields, not literal text - the whole reason doc-footer moved out of
	the body (see components.py's "doc-footer" comment) was so authors never
	again hand-write a page count that's wrong the moment pagination
	changes.

	python-docx has no field API, so this is raw OOXML: a field is a
	begin/instrText/separate/end fldChar sequence split across runs, per the
	OOXML spec (ECMA-376 17.16.18) - there is no simpler supported shape.
	"""
	paragraph.add_run("   |   Page ")
	_add_field_run(paragraph, "PAGE")
	paragraph.add_run(" of ")
	_add_field_run(paragraph, "NUMPAGES")


def _add_field_run(paragraph, field_code: str) -> None:
	run = paragraph.add_run()
	r = run._r

	begin = OxmlElement("w:fldChar")
	begin.set(qn("w:fldCharType"), "begin")

	instr_text = OxmlElement("w:instrText")
	instr_text.set(qn("xml:space"), "preserve")
	instr_text.text = f" {field_code} "

	separate = OxmlElement("w:fldChar")
	separate.set(qn("w:fldCharType"), "separate")

	end = OxmlElement("w:fldChar")
	end.set(qn("w:fldCharType"), "end")

	r.append(begin)
	r.append(instr_text)
	r.append(separate)
	r.append(end)


def _docx_hex(value: str) -> str:
	"""Strip a leading ``#`` for OOXML consumers. ``THEME`` (components.py)
	stores colours as ``#RRGGBB`` to match CSS, but ``w:fill``/``w:color``
	XML attributes and ``RGBColor.from_string`` both expect bare 6-digit hex
	with no ``#`` - ``RGBColor.from_string("#2C5AA8")`` raises outright
	(confirmed: ``ValueError: invalid literal for int() with base 16: '#2'``
	on the leading pair). Every colour reaching this module now comes from
	``resolve_theme_token``, which returns ``THEME``'s ``#``-prefixed form
	verbatim, so this normalisation has to sit centrally rather than at each
	call site.
	"""
	return value.lstrip("#")


def _rgb_color(hex_value: str) -> RGBColor:
	return RGBColor.from_string(_docx_hex(hex_value))


def _shade_cell(cell, hex_color: str) -> None:
	"""Shade a table cell's background via a <w:shd> element on its tcPr -
	python-docx has no public cell-shading API.
	"""
	if not hasattr(cell, "_tc"):
		return
	tc_pr = cell._tc.get_or_add_tcPr()
	shd = OxmlElement("w:shd")
	shd.set(qn("w:val"), "clear")
	shd.set(qn("w:color"), "auto")
	shd.set(qn("w:fill"), _docx_hex(hex_color))
	tc_pr.append(shd)


def _shade_run(run, hex_color: str) -> None:
	"""Shade a single run's background via a <w:shd> element on its rPr -
	OOXML supports run-level shading (unlike cell shading, python-docx has
	no wrapper for this either), which is what gives the "status-badge"
	recipe its filled-pill look in Word. Tested in Word/LibreOffice: a
	<w:shd> on rPr renders as a solid highlight behind just that run's
	text, distinct from (and independent of) the paragraph's own shading.
	"""
	r_pr = run._r.get_or_add_rPr()
	shd = OxmlElement("w:shd")
	shd.set(qn("w:val"), "clear")
	shd.set(qn("w:color"), "auto")
	shd.set(qn("w:fill"), _docx_hex(hex_color))
	r_pr.append(shd)


def _set_table_borders(table, visible: bool) -> None:
	"""Turn a table's borders on (the built-in "Table Grid" style) or fully
	off via an explicit <w:tblBorders> with val="nil" on every edge -
	python-docx's default table has no guaranteed borderless style name, so
	"off" is expressed directly in XML rather than relying on one.
	"""
	if visible:
		table.style = "Table Grid"
		return

	tbl_pr = table._tbl.tblPr
	borders = OxmlElement("w:tblBorders")
	for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
		edge_el = OxmlElement(f"w:{edge}")
		edge_el.set(qn("w:val"), "nil")
		borders.append(edge_el)
	tbl_pr.append(borders)


def _set_column_widths(table, ratios: list[float]) -> None:
	"""Apply proportional column widths (e.g. split's [0.66, 0.34]) against
	a fixed 6" reference width. Word requires the width to be set on both
	the column AND every cell in it for a fixed-width table to render
	correctly - setting only one is a well-known python-docx gotcha.
	"""
	table.autofit = False
	reference_width = 6.0
	for index, ratio in enumerate(ratios):
		width = Inches(reference_width * ratio)
		if index < len(table.columns):
			table.columns[index].width = width
		for row in table.rows:
			if index < len(row.cells):
				row.cells[index].width = width


def _clear_cell(cell) -> None:
	"""Remove a freshly-created table cell's single default empty paragraph
	so subsequent add_paragraph()/add_table() calls don't leave a stray
	blank line before the real content. Only ever called immediately before
	adding real content, so the cell is never left without any block child
	(which OOXML requires).
	"""
	if not cell.paragraphs:
		return
	paragraph = cell.paragraphs[0]
	if len(cell.paragraphs) == 1 and not paragraph.runs and not paragraph.text:
		element = paragraph._p
		element.getparent().remove(element)


def html_to_docx(html: str) -> bytes:
	"""Convert the HTML produced by ``render_document_html()`` to a .docx file.

	Args:
		html: A full HTML document string as produced by
			``huf.ai.artifacts.render.html.render_document_html``.

	Returns:
		Raw bytes of a valid .docx file.
	"""
	builder = _BodyTreeBuilder()
	builder.feed(html)
	builder.close()

	# The document's effective palette: the registry defaults, updated with
	# whatever the author validly overrode in a <style> block's :root rule
	# (see _extract_theme). Computed once and threaded through every
	# _render_node()/recipe call below rather than re-parsed per component,
	# since the whole document shares one theme.
	theme = _extract_theme("".join(builder.style_text_parts))

	# The parser walks the whole document (including <head>/<style>), but
	# only tags from the sanitized vocabulary open nodes, so non-body content
	# (title, style rules) never produces nodes under root.
	doc = Document()

	for node in builder.root.children:
		_render_node(doc, node, theme)

	buffer = io.BytesIO()
	doc.save(buffer)
	return buffer.getvalue()
